from docx import Document
from io import BytesIO

def create_docx(text):
    doc = Document()
    doc.add_heading('Отчет AI Freelance Agent', 0)
    doc.add_paragraph(text)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio